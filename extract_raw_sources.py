#!/usr/bin/env python3
"""Extract raw text from benchmark source files under core/, grouped by S### source_id."""

from __future__ import annotations

import argparse
import csv
import hashlib
import logging
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".xlsx", ".xls", ".csv", ".docx", ".txt"}
SOURCE_ID_PATTERN = re.compile(r"^(S\d{3})_", re.IGNORECASE)


@dataclass
class FileRecord:
    source_id: str
    path: Path
    rel_path: str
    extension: str
    file_size: int
    file_sha256: str
    status: str  # ok | failed
    error: str = ""
    text: str = ""
    text_length: int = 0
    text_sha256: str = ""
    duplicate_file_group: str = ""
    source_has_multiple_files: bool = False


def setup_logging(log_path: Path) -> logging.Logger:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logger = logging.getLogger("extract_raw_sources")
    logger.setLevel(logging.DEBUG)
    logger.handlers.clear()

    fmt = logging.Formatter("%(asctime)s %(levelname)s %(message)s")

    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    logger.addHandler(fh)

    sh = logging.StreamHandler(sys.stdout)
    sh.setLevel(logging.INFO)
    sh.setFormatter(fmt)
    logger.addHandler(sh)

    return logger


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def parse_source_id(filename: str) -> str | None:
    match = SOURCE_ID_PATTERN.match(filename)
    if not match:
        return None
    return match.group(1).upper()


def read_bytes(path: Path) -> bytes:
    return path.read_bytes()


def extract_txt(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_csv(path: Path) -> str:
    import csv as csv_module

    lines: list[str] = []
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with path.open("r", encoding=encoding, newline="") as f:
                reader = csv_module.reader(f)
                for row in reader:
                    lines.append("\t".join(row))
            return "\n".join(lines)
        except UnicodeDecodeError:
            lines.clear()
            continue
    with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
        reader = csv_module.reader(f)
        for row in reader:
            lines.append("\t".join(row))
    return "\n".join(lines)


def extract_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    raw = extract_txt(path)
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            parts.append(f"[page {i + 1} extraction error: {exc}]")
            continue
        parts.append(page_text)
    return "\n\n".join(parts)


def extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def extract_xls(path: Path) -> str:
    import xlrd

    book = xlrd.open_workbook(str(path))
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"=== Sheet: {sheet.name} ===")
        for row_idx in range(sheet.nrows):
            row = sheet.row_values(row_idx)
            cells = ["" if v is None or v == "" else str(v) for v in row]
            if any(str(c).strip() for c in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


EXTRACTORS = {
    ".txt": extract_txt,
    ".csv": extract_csv,
    ".html": extract_html,
    ".htm": extract_html,
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xls,
}


def extract_text(path: Path, extension: str) -> str:
    extractor = EXTRACTORS.get(extension)
    if extractor is None:
        raise ValueError(f"unsupported extension: {extension}")
    return extractor(path)


def discover_files(core_dir: Path, logger: logging.Logger) -> list[Path]:
    found: list[Path] = []
    skipped_unsupported_ext = 0
    skipped_bad_name = 0

    for path in sorted(core_dir.rglob("*")):
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        source_id = parse_source_id(path.name)
        if source_id is None:
            skipped_bad_name += 1
            logger.debug("skip (name does not match S###_): %s", path)
            continue
        found.append(path)

    logger.info(
        "discovered %d source files (skipped %d unsupported-name, scanned all supported ext under core/)",
        len(found),
        skipped_bad_name,
    )
    if skipped_unsupported_ext:
        logger.debug("skipped %d files with unsupported extension", skipped_unsupported_ext)
    return found


def process_file(path: Path, core_dir: Path, logger: logging.Logger) -> FileRecord:
    source_id = parse_source_id(path.name) or "UNKNOWN"
    rel_path = str(path.relative_to(core_dir.parent))
    ext = path.suffix.lower()

    try:
        raw_bytes = read_bytes(path)
        file_sha = sha256_bytes(raw_bytes)
        file_size = len(raw_bytes)
    except Exception as exc:
        logger.error("failed to read file %s: %s", path, exc)
        return FileRecord(
            source_id=source_id,
            path=path,
            rel_path=rel_path,
            extension=ext.lstrip("."),
            file_size=0,
            file_sha256="",
            status="failed",
            error=f"read error: {exc}",
        )

    try:
        text = extract_text(path, ext)
        text = text.strip()
        text_len = len(text)
        text_sha = sha256_text(text) if text_len else ""
        if text_len == 0:
            logger.warning("empty text extracted from %s", path)
        return FileRecord(
            source_id=source_id,
            path=path,
            rel_path=rel_path,
            extension=ext.lstrip("."),
            file_size=file_size,
            file_sha256=file_sha,
            status="ok",
            text=text,
            text_length=text_len,
            text_sha256=text_sha,
        )
    except Exception as exc:
        logger.error("extraction failed for %s: %s", path, exc, exc_info=True)
        return FileRecord(
            source_id=source_id,
            path=path,
            rel_path=rel_path,
            extension=ext.lstrip("."),
            file_size=file_size,
            file_sha256=file_sha,
            status="failed",
            error=str(exc),
        )


def mark_duplicates(records: list[FileRecord], logger: logging.Logger) -> None:
    by_file_hash: dict[str, list[FileRecord]] = defaultdict(list)
    by_text_hash: dict[str, list[FileRecord]] = defaultdict(list)
    by_source: dict[str, list[FileRecord]] = defaultdict(list)

    for rec in records:
        if rec.file_sha256:
            by_file_hash[rec.file_sha256].append(rec)
        if rec.text_sha256:
            by_text_hash[rec.text_sha256].append(rec)
        by_source[rec.source_id].append(rec)

    for sha, group in by_file_hash.items():
        if len(group) < 2:
            continue
        group_id = f"file_sha256:{sha[:16]}"
        paths = [r.rel_path for r in group]
        logger.warning("duplicate file bytes (%d files): %s", len(group), paths)
        for rec in group:
            rec.duplicate_file_group = group_id

    multi_source = {sid: grp for sid, grp in by_source.items() if len(grp) > 1}
    for sid, group in sorted(multi_source.items()):
        paths = [r.rel_path for r in group]
        logger.warning("source_id %s has %d files: %s", sid, len(group), paths)
        for rec in group:
            rec.source_has_multiple_files = True

    for sha, group in by_text_hash.items():
        if len(group) < 2:
            continue
        # Only log if not already same file hash group
        unique_file_hashes = {r.file_sha256 for r in group}
        if len(unique_file_hashes) > 1:
            paths = [r.rel_path for r in group]
            logger.warning(
                "duplicate extracted text (%d files, different bytes): %s",
                len(group),
                paths,
            )


def write_grouped_raw_text(
    records: list[FileRecord],
    raw_text_dir: Path,
    logger: logging.Logger,
) -> dict[str, Path]:
    by_source: dict[str, list[FileRecord]] = defaultdict(list)
    for rec in records:
        if rec.status == "ok":
            by_source[rec.source_id].append(rec)

    written: dict[str, Path] = {}
    for source_id in sorted(by_source):
        group = sorted(by_source[source_id], key=lambda r: r.rel_path)
        chunks: list[str] = []
        for rec in group:
            header = f"===== FILE: {rec.rel_path} ====="
            chunks.append(f"{header}\n{rec.text}")
        combined = "\n\n".join(chunks) + "\n"
        out_path = raw_text_dir / f"{source_id}.txt"
        out_path.write_text(combined, encoding="utf-8")
        written[source_id] = out_path
        logger.info("wrote %s (%d files, %d chars)", out_path, len(group), len(combined))

    failed_sources = {r.source_id for r in records if r.status == "failed"}
    ok_sources = set(by_source)
    for sid in sorted(failed_sources - ok_sources):
        logger.error("no raw_text output for %s (all files failed)", sid)

    return written


def write_index_csv(records: list[FileRecord], index_path: Path) -> None:
    fieldnames = [
        "source_id",
        "file_path",
        "file_name",
        "format",
        "status",
        "error_message",
        "file_size_bytes",
        "file_sha256",
        "text_length",
        "text_sha256",
        "duplicate_file_group",
        "source_has_multiple_files",
        "raw_text_output",
        "processed_at_utc",
    ]
    timestamp = datetime.now(timezone.utc).isoformat()

    source_outputs: dict[str, str] = {}
    for rec in records:
        if rec.status == "ok":
            source_outputs[rec.source_id] = f"raw_text/{rec.source_id}.txt"

    with index_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for rec in sorted(records, key=lambda r: (r.source_id, r.rel_path)):
            writer.writerow(
                {
                    "source_id": rec.source_id,
                    "file_path": rec.rel_path,
                    "file_name": rec.path.name,
                    "format": rec.extension,
                    "status": rec.status,
                    "error_message": rec.error,
                    "file_size_bytes": rec.file_size,
                    "file_sha256": rec.file_sha256,
                    "text_length": rec.text_length,
                    "text_sha256": rec.text_sha256,
                    "duplicate_file_group": rec.duplicate_file_group,
                    "source_has_multiple_files": rec.source_has_multiple_files,
                    "raw_text_output": source_outputs.get(rec.source_id, ""),
                    "processed_at_utc": timestamp,
                }
            )


def run(core_dir: Path, output_dir: Path) -> int:
    core_dir = core_dir.resolve()
    output_dir = output_dir.resolve()
    raw_text_dir = output_dir / "raw_text"
    raw_text_dir.mkdir(parents=True, exist_ok=True)

    log_path = output_dir / "extraction.log"
    index_path = output_dir / "raw_sources_index.csv"
    logger = setup_logging(log_path)

    logger.info("core_dir=%s output_dir=%s", core_dir, output_dir)

    if not core_dir.is_dir():
        logger.error("core directory does not exist: %s", core_dir)
        return 1

    paths = discover_files(core_dir, logger)
    if not paths:
        logger.error("no matching source files found under %s", core_dir)
        write_index_csv([], index_path)
        return 1

    records: list[FileRecord] = []
    for path in paths:
        logger.info("processing %s", path)
        records.append(process_file(path, core_dir, logger))

    mark_duplicates(records, logger)
    write_grouped_raw_text(records, raw_text_dir, logger)
    write_index_csv(records, index_path)

    ok = sum(1 for r in records if r.status == "ok")
    failed = sum(1 for r in records if r.status == "failed")
    sources = len({r.source_id for r in records})
    sources_written = len({r.source_id for r in records if r.status == "ok"})

    logger.info(
        "done: %d files, %d ok, %d failed, %d source_ids, %d raw_text files",
        len(records),
        ok,
        failed,
        sources,
        sources_written,
    )
    logger.info("index: %s", index_path)
    logger.info("log: %s", log_path)

    return 0 if failed == 0 else 2


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--core-dir",
        type=Path,
        default=Path(__file__).resolve().parent / "core",
        help="Folder containing S###_* source files (default: ./core)",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path(__file__).resolve().parent,
        help="Directory for raw_text/, index CSV, and log (default: project root)",
    )
    args = parser.parse_args()
    sys.exit(run(args.core_dir, args.output_dir))


if __name__ == "__main__":
    main()
